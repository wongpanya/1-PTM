from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT


OUTPUT = r"C:\Users\Wongpanya.Nu\Documents\1-PTM\ODOS_Prototype_Implementation_Plan.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = "Tahoma"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)


def set_table_borders(table, color="D9E2EC"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_run(run, size=None, bold=None, color=None):
    run.font.name = "Tahoma"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(text)
    style_run(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True, color="1F4D78" if level < 3 else "0B2545")
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        style_run(r, size=11, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        style_run(r2, size=11)
    else:
        r = p.add_run(text)
        style_run(r, size=11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.1)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    style_run(r, size=10.5)


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.1)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    style_run(r, size=10.5)


def add_matrix(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    if widths:
        for i, w in enumerate(widths):
            table.columns[i].width = Inches(w)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], "E8EEF5")
        set_cell_text(hdr[i], h, bold=True, color="0B2545", size=9.5)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=9)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, "B7C9DD")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F8FC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    style_run(r, size=11, bold=True, color="1F4D78")
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    style_run(r2, size=10.5)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

styles = doc.styles
styles["Normal"].font.name = "Tahoma"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
styles["Normal"].font.size = Pt(11)

for style_name, size, color in [
    ("Heading 1", 16, "1F4D78"),
    ("Heading 2", 13, "2E74B5"),
    ("Heading 3", 12, "0B2545"),
]:
    st = styles[style_name]
    st.font.name = "Tahoma"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(12)
    st.paragraph_format.space_after = Pt(6)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("แผนการดำเนินงาน Prototype\nODOS Policy Analytics Platform")
style_run(r, size=20, bold=True, color="0B2545")
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("ทบทวนกรอบงาน ทรัพยากร และแนวทางพัฒนาระบบภายใน 2-3 วัน")
style_run(r, size=12, color="4A5568")

add_callout(
    doc,
    "วัตถุประสงค์ของเอกสาร",
    "สรุปแนวทางดำเนินงานเพื่อพัฒนา Prototype จากข้อมูลในไฟล์โครงการที่มีอยู่ ให้เห็นภาพระบบวิเคราะห์ข้อมูลผู้รับทุน ระบบกรอกข้อมูลเสริมรายปี Dashboard โมดูลพยากรณ์เบื้องต้น และข้อเสนอเชิงนโยบายสำหรับนำไปของบสนับสนุนในระยะใช้งานจริง",
)

add_heading(doc, "1. หลักคิดของ Prototype", 1)
add_body(doc, "Prototype นี้ควรออกแบบเพื่อพิสูจน์แนวคิด ไม่ใช่ระบบ production เต็มรูปแบบ เป้าหมายคือทำให้ผู้บริหารเห็นว่าข้อมูลผู้รับทุนที่มีอยู่สามารถแปลงเป็นระบบสนับสนุนนโยบายได้จริง และเห็นช่องว่างข้อมูลที่ต้องลงทุนเพิ่มเติมในระยะถัดไป")
for item in [
    "ใช้ข้อมูลเท่าที่มีใน project นี้เป็นฐานหลัก โดยเฉพาะไฟล์ฐานข้อมูลผู้รับทุน ODOS",
    "แยก Core Scholarship Database ออกจาก Annual External Indicators เพื่อรองรับข้อมูลภายนอกที่เปลี่ยนรายปี",
    "ทำระบบให้ deploy บนเว็บได้ง่าย ใช้เครื่องมือฟรีหรือ free tier ก่อน",
    "ออกแบบ agents เป็นโมดูลการทำงานในระบบก่อน แล้วค่อยต่อยอดเป็น AI/automation ในระยะจริง",
    "จัดการกฎหมายและธรรมาภิบาลข้อมูลในระดับ Prototype ด้วยการซ่อนข้อมูลส่วนบุคคล แสดงผลแบบ aggregate และบันทึก log เบื้องต้น",
]:
    add_bullet(doc, item)

add_heading(doc, "2. โครงสร้างระบบที่เสนอ", 1)
add_matrix(
    doc,
    ["ส่วนระบบ", "หน้าที่หลัก", "ขอบเขต Prototype"],
    [
        ["Core Scholarship Database", "จัดเก็บและจัดระเบียบข้อมูลผู้รับทุน รุ่น พื้นที่ สาขา ประเทศ สถานะ อาชีพ และรายได้", "นำเข้า Excel, clean เฉพาะ field สำคัญ, สร้าง dataset กลาง"],
        ["Annual External Indicators", "รองรับข้อมูลต้นทุนทุน ตลาดแรงงาน GDP รายได้เฉลี่ย SDGs และ priority policy รายปี", "ทำ template/form กรอกข้อมูลและตัวอย่าง mock data"],
        ["Dashboard", "แสดงภาพรวมและ drill-down สำหรับผู้บริหารและเจ้าหน้าที่", "Overview, Data Quality, Analytics, Risk, Policy"],
        ["Risk & Forecast", "คำนวณคะแนนความเสี่ยงและพยากรณ์ตั้งต้น", "rule-based score หรือ model เบื้องต้นที่อธิบายได้"],
        ["Policy Recommendation", "แปลงผลวิเคราะห์เป็นข้อเสนอเชิงนโยบาย", "ranking สาขา/พื้นที่/กลุ่มเสี่ยง พร้อมเหตุผล"],
        ["Governance", "ควบคุมการใช้ข้อมูลและลดความเสี่ยง PDPA", "mask PII, role mockup, audit log เบื้องต้น"],
    ],
    widths=[1.8, 2.9, 2.8],
)

add_heading(doc, "3. แผนดำเนินงานรายส่วน", 1)
add_heading(doc, "3.1 ข้อมูลและฐานข้อมูลกลาง", 2)
for item in [
    "ตรวจสอบโครงสร้างชีต คอลัมน์ จำนวนข้อมูล และคุณภาพข้อมูล",
    "กำหนด field หลักสำหรับ Prototype เช่น ID, รุ่น, จังหวัด, อำเภอ, เพศ, ประเทศ, สาขา, สถานะ, อาชีพ, รายได้ และความสอดคล้องงาน",
    "ทำ data cleaning เฉพาะจุดที่มีผลต่อ dashboard และโมเดล เช่น N/A, ค่าว่าง, #NUM!, ชื่อจังหวัด/ประเทศที่สะกดต่างกัน",
    "สร้าง data dictionary จากชีต Remark เพื่อใช้เป็นรายการมาตรฐาน",
    "แปลงข้อมูลเป็นตารางกลางสำหรับระบบ เช่น students, education, employment, status, geography",
]:
    add_bullet(doc, item)

add_heading(doc, "3.2 ระบบข้อมูลเสริมรายปี", 2)
for item in [
    "ออกแบบ template สำหรับกรอกข้อมูลภายนอกที่เปลี่ยนตามปี",
    "รองรับมิติปีงบประมาณ แหล่งข้อมูล วันที่อัปเดต ระดับข้อมูล และค่าตัวชี้วัด",
    "แยกประเภทข้อมูล เช่น cost, labor_market, income_benchmark, provincial_indicator, SDGs, policy_priority",
    "Prototype ใช้ mock/sample data เพื่อแสดงว่าในอนาคตสามารถเติมข้อมูลและนำไปคำนวณ ROI/SROI หรือ workforce demand ได้",
]:
    add_bullet(doc, item)

add_heading(doc, "3.3 Dashboard และ Analytics", 2)
for item in [
    "สร้าง Overview Dashboard: จำนวนผู้รับทุนตามรุ่น พื้นที่ ประเทศ สาขา และสถานะ",
    "สร้าง Data Quality Dashboard: ความครบถ้วน ค่าว่าง ข้อมูลผิดรูปแบบ และ field ที่พร้อมใช้",
    "สร้าง Analytics Dashboard: completion, dropout, occupation, income, field-job fit, local fit",
    "เพิ่ม filter หลัก เช่น รุ่น จังหวัด ภูมิภาค ประเทศ สาขา และสถานะ",
]:
    add_bullet(doc, item)

add_heading(doc, "3.4 Risk, Forecast และ Policy", 2)
for item in [
    "เริ่มจาก risk score แบบอธิบายได้ เช่น สถานะไม่สำเร็จ, ติดตามข้อมูลไม่ได้, ทำงานไม่ตรงสาขา, ไม่สอดคล้องท้องถิ่น",
    "ทำโมดูล Graduation Success และ Scholarship Risk ก่อน เพราะข้อมูลพร้อมที่สุด",
    "ทำ Field Recommendation และ Area-based Allocation เบื้องต้นจาก completion, employment, fit และพื้นที่",
    "ทำ Policy Recommendation Page ที่สรุปข้อเสนอพร้อมเหตุผล ไม่ให้เป็น black box",
]:
    add_bullet(doc, item)

add_heading(doc, "4. Agents ที่ควรออกแบบใน Prototype", 1)
add_matrix(
    doc,
    ["Agent / Module", "บทบาท", "ผลลัพธ์ใน Prototype"],
    [
        ["Data Steward Agent", "ตรวจคุณภาพข้อมูลและจัดมาตรฐาน", "รายงาน data quality, รายการ field ที่ต้องแก้, completeness score"],
        ["Analytics Agent", "สรุปภาพรวมและวิเคราะห์เชิงพรรณนา/วินิจฉัย", "chart, pivot, insight ตามรุ่น พื้นที่ สาขา ประเทศ"],
        ["Risk Prediction Agent", "คำนวณความเสี่ยงของผู้รับทุนหรือกลุ่มข้อมูล", "risk score, กลุ่มเสี่ยง, เหตุผลของคะแนน"],
        ["Policy Recommendation Agent", "แปลงผลวิเคราะห์เป็นข้อเสนอเชิงนโยบาย", "ข้อเสนอพื้นที่/สาขา/กลุ่มติดตาม พร้อมหลักฐาน"],
        ["External Indicator Agent", "จัดการข้อมูลเสริมรายปี", "template กรอกข้อมูลและ readiness score สำหรับ ROI/SROI"],
        ["Governance Agent", "ควบคุมการเปิดเผยข้อมูลและการใช้งาน", "mask PII, role mockup, audit log, export policy"],
    ],
    widths=[1.7, 2.6, 3.0],
)

add_heading(doc, "5. เทคโนโลยีที่แนะนำสำหรับ Deploy ฟรีก่อน", 1)
add_matrix(
    doc,
    ["องค์ประกอบ", "ตัวเลือกแนะนำ", "เหตุผล"],
    [
        ["Web App", "Streamlit", "ทำ dashboard, filter, upload, chart และ deploy ได้เร็ว เหมาะกับ Prototype 2-3 วัน"],
        ["Data Processing", "Python + Pandas", "อ่าน Excel, clean data, สร้าง metrics และ scoring ได้รวดเร็ว"],
        ["Visualization", "Plotly", "ทำ interactive charts และ map/treemap/bar chart ได้ง่าย"],
        ["Database Prototype", "SQLite หรือ CSV/Parquet", "ไม่ต้องตั้ง server ซับซ้อน เหมาะกับ demo"],
        ["Deployment", "Streamlit Community Cloud หรือ Hugging Face Spaces", "มี free tier และเหมาะกับ demo dashboard"],
        ["Version Control", "GitHub", "ใช้เก็บ source code และเชื่อม deployment ได้สะดวก"],
    ],
    widths=[1.5, 2.2, 3.4],
)

add_heading(doc, "6. Governance และ PDPA ในระดับ Prototype", 1)
add_body(doc, "ประเด็นกฎหมายและธรรมาภิบาลข้อมูลสามารถออกแบบให้จัดการในระบบได้ แต่ใน Prototype ควรทำเฉพาะมาตรการพื้นฐานที่เห็นภาพและลดความเสี่ยงก่อน")
for item in [
    "ซ่อนหรือไม่แสดงข้อมูลส่วนบุคคล เช่น ชื่อ เบอร์โทร ที่อยู่ เลขที่สัญญา และรายละเอียดติดต่อ",
    "แสดงผลเป็นข้อมูล aggregate เป็นหลัก เช่น จำนวน อัตรา ค่าเฉลี่ย และ ranking กลุ่ม",
    "ทำ role mockup เช่น Admin, Analyst, Viewer เพื่อสื่อแนวคิดสิทธิ์การเข้าถึง",
    "บันทึก audit log เบื้องต้น เช่น วันที่ import, จำนวน records, ผู้ใช้งาน และกิจกรรม export",
    "ใส่ data use notice ว่าข้อมูลใช้เพื่อวิเคราะห์เชิงนโยบายและยังเป็น Prototype",
]:
    add_bullet(doc, item)

add_heading(doc, "7. Timeline 2-3 วัน", 1)
add_matrix(
    doc,
    ["วัน", "งานหลัก", "ผลลัพธ์"],
    [
        ["วันที่ 1", "ตรวจข้อมูล, clean field สำคัญ, สร้าง data dictionary, ออกแบบ schema และ metrics", "cleaned dataset, data quality summary, schema prototype"],
        ["วันที่ 2", "พัฒนา web dashboard หน้า Overview, Data Quality และ Analytics", "dashboard ใช้งานได้พร้อม filter และ chart หลัก"],
        ["วันที่ 3", "เพิ่ม Risk Score, Policy Recommendation, External Indicators template, Governance mockup และเตรียม deploy", "Prototype พร้อมนำเสนอและเว็บ demo"],
    ],
    widths=[1.0, 3.6, 2.7],
)

add_heading(doc, "8. Deliverables สำหรับนำเสนอของบสนับสนุน", 1)
for item in [
    "เว็บ Prototype สำหรับสาธิตการวิเคราะห์ข้อมูลผู้รับทุน",
    "Dashboard ภาพรวมและ dashboard คุณภาพข้อมูล",
    "โมดูล risk score และ policy recommendation เบื้องต้น",
    "Template สำหรับกรอกข้อมูลเสริมรายปี",
    "เอกสาร architecture และ roadmap สำหรับพัฒนา production",
    "รายการข้อมูลที่ต้องจัดหาเพิ่มเพื่อทำ ROI/SROI, workforce demand และ national impact",
    "ข้อเสนอทรัพยากร บุคลากร งบประมาณ และระยะเวลาดำเนินงานระยะจริง",
]:
    add_bullet(doc, item)

add_heading(doc, "9. ขอบเขตที่ควรชะลอไว้หลัง Prototype", 1)
for item in [
    "Machine learning ซับซ้อนหรือโมเดล black-box ที่ยังอธิบายไม่ได้",
    "การเชื่อมข้อมูลภายนอกอัตโนมัติจากหลายหน่วยงาน",
    "ระบบ login/permission ระดับองค์กรเต็มรูปแบบ",
    "ROI/SROI สมบูรณ์ในระดับรายบุคคล หากยังไม่มีต้นทุนทุนและตัวชี้วัด social value ที่ครบ",
    "ระบบ production ที่ต้องรองรับ concurrency, backup, security hardening และ SLA",
]:
    add_bullet(doc, item)

add_callout(
    doc,
    "ข้อเสนอแนะสรุป",
    "เริ่มจาก Prototype ที่ทำให้เห็นคุณค่าของข้อมูลจริงก่อน: clean data, dashboard, risk score, policy recommendation และ annual indicator template จากนั้นใช้ผลลัพธ์นี้เป็นหลักฐานเชิงประจักษ์เพื่อเสนอของบพัฒนาระบบ production ที่ครบถ้วนต่อไป",
)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("ODOS Policy Analytics Prototype | Draft for Policy Proposal")
style_run(r, size=8.5, color="667085")

doc.save(OUTPUT)
print(OUTPUT)
