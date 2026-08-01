from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "ODOS_Project_Framework_and_Scope_Full_TH.docx"

FONT = "TH Sarabun New"
BLUE = RGBColor(31, 78, 121)
ACCENT = RGBColor(46, 116, 181)
DARK = RGBColor(31, 31, 31)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F3F6F8"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:cs"), FONT)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, fill=None, align=None):
    if fill:
        shade_cell(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, 14, bold, DARK)


def add_para(doc, text="", style=None, size=15, bold=False, color=DARK, after=6, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size, bold, color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, 14.5, False, DARK)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, 14.5, False, DARK)


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        set_run_font(run, 20 if level == 1 else 17 if level == 2 else 15.5, True, BLUE if level < 3 else ACCENT)
    if not p.runs:
        r = p.add_run(text)
        set_run_font(r, 20 if level == 1 else 17, True, BLUE)
    else:
        p.runs[0].text = text
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, fill=LIGHT_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, 15, True, BLUE)
    p.add_run("\n")
    r2 = p.add_run(body)
    set_run_font(r2, 14.5, False, DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_framework_overview(doc):
    add_heading(doc, "ภาพรวมกรอบงานระดับระบบ", 2)
    add_table(
        doc,
        ["Input", "Data Management", "Analytics", "Decision Support", "Governance"],
        [
            [
                "ข้อมูลผู้รับทุน การศึกษา การใช้ทุน ผลลัพธ์หลังเรียน และข้อมูลเสริมรายปี",
                "นำเข้า ตรวจ schema ทำความสะอาด ทำมาตรฐาน จัดเก็บ SQLite/CSV และบันทึก lineage",
                "Overview, Data Quality, Analytics, Risk Forecast, Policy Recommendation และ External Indicators",
                "ข้อเสนอเชิงนโยบายแบบอธิบายได้ ปรับน้ำหนักได้ และต้องผ่าน human review",
                "aggregate-first, no-direct-PII, role mockup, audit/export log, minimum group-size suppression",
            ]
        ],
        [1.25, 1.35, 1.35, 1.3, 1.25],
    )


def setup_styles(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    sec.header_distance = Inches(0.4)
    sec.footer_distance = Inches(0.35)

    styles = doc.styles
    for name in ["Normal", "List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)
        style.font.size = Pt(15)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 20, BLUE),
        ("Heading 2", 17, BLUE),
        ("Heading 3", 15.5, ACCENT),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)


def build_doc():
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run("กรอบงานและขอบเขตโครงการฉบับเต็ม")
    set_run_font(r, 28, True, BLUE)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("ODOS Policy Analytics Prototype")
    set_run_font(r, 22, True, ACCENT)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("ฉบับปรับปรุงตามสถานะระบบล่าสุด | 29 กรกฎาคม 2026")
    set_run_font(r, 14, False, MUTED)

    add_framework_overview(doc)

    add_callout(
        doc,
        "สาระสำคัญ",
        "โครงการนี้เป็น Prototype สำหรับพิสูจน์ว่าข้อมูลผู้รับทุน 1 อำเภอ 1 ทุน สามารถพัฒนาเป็นระบบสนับสนุนการวิเคราะห์และข้อเสนอเชิงนโยบายได้จริง โดยเน้น dashboard, data quality, analytics, risk/forecast, policy recommendation, external indicators, governance และ help/documentation ภายใต้หลัก aggregate-first, no-direct-PII และไม่ใช้เป็นระบบตัดสินใจจัดสรรทุนอัตโนมัติ",
    )

    doc.add_page_break()
    add_heading(doc, "1. บทนำและวัตถุประสงค์", 1)
    add_para(doc, "ODOS Policy Analytics Prototype ถูกออกแบบเพื่อเปลี่ยนข้อมูลผู้รับทุน 1 อำเภอ 1 ทุน จากชุดข้อมูลเชิงทะเบียนให้กลายเป็นเครื่องมือวิเคราะห์เชิงนโยบายที่ผู้บริหาร นักวิเคราะห์ และผู้ดูแลข้อมูลสามารถใช้สำรวจสถานการณ์ ตรวจสอบคุณภาพข้อมูล ประเมินความเสี่ยง และจัดทำข้อเสนอเชิงนโยบายได้อย่างมีหลักฐานรองรับ")
    add_para(doc, "เป้าหมายของ Prototype ไม่ใช่การสร้างระบบ production ที่พร้อมใช้ทั่วทั้งองค์กรทันที แต่เป็นการพิสูจน์แนวคิด ระบบข้อมูล กระบวนการวิเคราะห์ และรูปแบบการกำกับดูแลข้อมูลในขอบเขตที่ควบคุมได้ โดยใช้ข้อมูลที่ผ่านการลดการระบุตัวตนและนำเสนอผลลัพธ์ระดับภาพรวมเป็นหลัก")
    add_bullet(doc, "พิสูจน์ความเป็นไปได้ของระบบสนับสนุนนโยบายจากข้อมูลทุน ODOS")
    add_bullet(doc, "สร้าง dashboard และ analytical workflow สำหรับการสำรวจผลลัพธ์การศึกษาและการทำงาน")
    add_bullet(doc, "วางรากฐาน data quality, risk scoring, policy recommendation และ governance")
    add_bullet(doc, "เตรียม roadmap สำหรับต่อยอดสู่ระบบ production เมื่อข้อมูล นิยาม กฎ และ governance ผ่านการรับรอง")

    add_heading(doc, "2. กรอบแนวคิดโครงการ", 1)
    add_para(doc, "กรอบงานเริ่มจากข้อมูลนำเข้า เช่น ข้อมูลผู้รับทุน ข้อมูลการศึกษา ข้อมูลการใช้ทุน ข้อมูลผลการศึกษา ข้อมูลการทำงาน และข้อมูลเสริมด้านสังคม จากนั้นจัดการข้อมูลผ่านการ integration, cleaning, standardization และจัดเก็บเป็นฐานข้อมูลกลาง ก่อนนำเข้าสู่การวิเคราะห์ 4 มิติ ได้แก่ descriptive, diagnostic, predictive และ prescriptive analytics")
    add_table(
        doc,
        ["องค์ประกอบ", "บทบาทในกรอบงาน"],
        [
            ["Input", "รวบรวมข้อมูลผู้รับทุน การศึกษา การใช้ทุน ผลการเรียน การทำงาน และบริบทเสริมที่เกี่ยวข้อง"],
            ["Data Management", "รวมข้อมูล ทำความสะอาด ทำมาตรฐาน และเตรียมฐานข้อมูลกลางสำหรับการวิเคราะห์"],
            ["Data Analytics", "วิเคราะห์ภาพรวม วินิจฉัยปัญหา ประเมินแนวโน้ม และเสนอทางเลือกเชิงนโยบาย"],
            ["Key Forecast", "ประเมินความต้องการกำลังคน สาขาอนาคต ROI/SROI โอกาสการมีงานทำ brain drain graduation success area allocation leadership national impact และ scholarship risk"],
            ["Policy & Decision Support", "แปลงผลวิเคราะห์เป็นข้อเสนอที่อธิบายได้ ปรับน้ำหนักได้ และให้ผู้เชี่ยวชาญใช้พิจารณา"],
            ["Help & Documentation", "ค้นหาคู่มือ FAQ และถามผู้ช่วย Local AI จากเอกสารโครงการ โดยไม่ส่งฐานข้อมูลผู้รับทุนหรือ PII"],
            ["Ultimate Outcome", "สนับสนุนกำลังคนคุณภาพสูง ลดความเหลื่อมล้ำ สร้างผลตอบแทนทางเศรษฐกิจและสังคม และสนับสนุน SDGs"],
        ],
        [1.5, 5.0],
    )

    add_heading(doc, "3. ขอบเขต Prototype", 1)
    add_para(doc, "ขอบเขตงานรอบ Prototype ครอบคลุมการสร้างระบบต้นแบบแบบ local/prototype โดยใช้ Streamlit, Python, Pandas, Plotly และ SQLite/CSV เป็นแกนหลัก ข้อมูลหลักมาจาก sample หรือ cleaned dataset ที่ไม่มี direct PII และเน้นการสาธิต workflow มากกว่าการให้บริการ production")
    add_bullet(doc, "ใช้ข้อมูลเท่าที่มีใน project เป็นฐานหลัก และไม่อ่าน raw Excel ทุกครั้งใน runtime ของ web app")
    add_bullet(doc, "ใช้ฐานข้อมูลกลางแบบ local/prototype เช่น SQLite และไฟล์ CSV ที่ผ่านการเตรียมข้อมูล")
    add_bullet(doc, "สร้าง web dashboard 8 หน้า ได้แก่ Overview, Data Quality, Analytics, Risk Forecast, Policy Recommendation, External Indicators, Governance และ Help & Documentation")
    add_bullet(doc, "พัฒนา risk score และ policy recommendation แบบอธิบายได้ โดยใช้ rule-based logic หรือ model เบื้องต้น")
    add_bullet(doc, "จัดทำ template สำหรับ external indicators รายปี และ governance mockup เช่น PII masking, role concept, audit/export log")
    add_bullet(doc, "รองรับการค้นหาเอกสารและ AI-assisted help ผ่าน Local AI/Ollama โดยจำกัดเฉพาะเอกสารโครงการ ไม่เชื่อมฐานข้อมูลรายบุคคล")
    add_bullet(doc, "รองรับการทดสอบ acceptance, privacy check, data validation, label validation และ handover สำหรับผู้ใช้งานทดลอง")

    add_heading(doc, "4. สิ่งที่อยู่นอกขอบเขต", 1)
    add_para(doc, "เพื่อให้ความคาดหวังชัดเจน Prototype ไม่ครอบคลุมงาน production ที่ต้องมี security, reliability, compliance และ operating model เต็มรูปแบบ")
    for item in [
        "ระบบ production สำหรับใช้งานจริงระดับองค์กร",
        "authentication/authorization แบบเต็มรูปแบบและ identity lifecycle",
        "security hardening, encryption, monitoring, backup, disaster recovery และ SLA",
        "database server กลางหรือ data warehouse production",
        "machine learning production pipeline ที่ผ่าน leakage review, calibration, monitoring และ formal production approval",
        "API เชื่อมโยงข้อมูลภายนอกอัตโนมัติ",
        "การตัดสินใจจัดสรรทุนอัตโนมัติ",
        "AI ภายนอกหรือ cloud AI ที่ส่งข้อมูลออกนอกเครื่องโดยไม่มี privacy/DPA review",
        "PDPA workflow เต็มรูปแบบและกระบวนการอนุมัติ/เก็บรักษาข้อมูลระดับ production",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. กลุ่มผู้ใช้หลัก", 1)
    add_table(
        doc,
        ["กลุ่มผู้ใช้", "ความต้องการหลัก", "ผลลัพธ์ที่ใช้"],
        [
            ["ผู้บริหารและผู้กำหนดนโยบาย", "เห็นภาพรวม ผลลัพธ์ ความเสี่ยง และข้อเสนอ", "KPI, dashboard, policy ranking, executive scenario"],
            ["นักวิเคราะห์ข้อมูล/นโยบาย", "สำรวจข้อมูล เปรียบเทียบมิติ วิเคราะห์เหตุผล", "filter, chart, metric definition, risk explanation"],
            ["เจ้าหน้าที่ดูแลข้อมูล", "ตรวจคุณภาพข้อมูล เตรียมชุดข้อมูล และติดตามข้อผิดพลาด", "validation issues, data dictionary, cleaning report, import log"],
            ["เจ้าของ governance/ข้อมูล", "ตรวจการใช้ข้อมูล บทบาท การ export และข้อจำกัด", "PII masking, audit/export log, aggregate-only policy"],
            ["ผู้ใช้ทดลอง/ผู้ส่งมอบงาน", "ค้นหาคู่มือ เรียนรู้ขั้นตอน และถามคำถามจากเอกสาร", "Help Documentation, FAQ, Local AI assistant"],
        ],
        [1.6, 2.4, 2.5],
    )

    add_heading(doc, "6. หน้าระบบและขอบเขตฟังก์ชัน", 1)
    add_table(
        doc,
        ["หน้า", "เป้าหมาย", "ตัวอย่างข้อมูล/ฟังก์ชัน"],
        [
            ["Overview", "แสดงภาพรวมผู้รับทุนและผลลัพธ์หลัก", "จำนวนผู้รับทุน cohort ภูมิภาค จังหวัด อำเภอ สถานะการศึกษา ประเทศ สาขา อาชีพหลังเรียน"],
            ["Data Quality", "แสดงความพร้อมข้อมูลและช่องว่างที่ต้องปรับปรุง", "completeness, missing value, error-like value, readiness, data dictionary"],
            ["Analytics", "วิเคราะห์เชิงพรรณนาและวินิจฉัย", "completion rate, dropout, employment, income, field-job fit, local fit"],
            ["Risk & Forecast", "สาธิต risk score และ graduation status", "คะแนนความเสี่ยง ระดับความเสี่ยง components rule version limitations"],
            ["Policy Recommendation", "จัดอันดับข้อเสนอเชิงนโยบายจากสูตรและน้ำหนัก", "field recommendation, area-based allocation, weight adjustment, evidence columns"],
            ["External Indicators", "รองรับข้อมูลภายนอกรายปี", "template ต้นทุน ตลาดแรงงาน รายได้ GDP/จังหวัด SDGs policy priority"],
            ["Governance", "แสดงแนวทางกำกับข้อมูลและ privacy", "role mockup, minimum group size, PII masking, audit log, export policy"],
            ["Help & Documentation", "ค้นหาเอกสาร FAQ และถาม Local AI จากเอกสารโครงการ", "documentation search, quick links, FAQ, Ollama localhost assistant, privacy warning"],
        ],
        [1.35, 2.2, 2.95],
    )

    add_heading(doc, "7. ข้อมูลนำเข้าและข้อมูลส่งออก", 1)
    add_para(doc, "ข้อมูลเริ่มต้นมาจาก workbook ภายในโครงการ เช่น 690724 DB_ODOS Students+.xlsx ที่ประกอบด้วยชีต DB_Students และ Remark โดย raw data ควรเก็บแยกจาก public repository และระบบใช้งานควรอ่านจาก cleaned CSV หรือ SQLite ที่สร้างจาก pipeline แล้ว")
    add_table(
        doc,
        ["ประเภท", "รายการ"],
        [
            ["Input หลัก", "DB_Students, Remark, no-PII sample, validation data, annual external indicators template, label review template"],
            ["Output Phase 1/4", "cleaned dataset, data dictionary, validation issues, field cleaning report, processing log, import manifest"],
            ["ฐานข้อมูล Prototype", "students, education_records, employment_records, scholarship_status, geography_reference, external_indicators, data_import_log, risk_scores, policy_recommendations, audit_logs"],
            ["นโยบายข้อมูล", "raw data ไม่แก้ไข, failed records ไม่ถูกลบ, issues ถูกบันทึก, direct PII ไม่แสดงบน dashboard และไม่ export; Help/AI ไม่เชื่อมฐานข้อมูลผู้รับทุน"],
        ],
        [1.7, 4.8],
    )

    add_heading(doc, "8. สถาปัตยกรรมและเทคโนโลยี", 1)
    add_para(doc, "สถาปัตยกรรมแบ่งเป็นชั้นข้อมูล config business logic UI และ validation เพื่อให้ Prototype ดูแลได้ง่ายและต่อยอดได้เป็น production ในอนาคต โครงสร้างปัจจุบันใช้ data/sample และ data/reference สำหรับข้อมูล, config/*.yaml สำหรับกฎและนิยาม, src/* สำหรับ logic, app.py/pages/*.py สำหรับ UI และ scripts/tests สำหรับ validation")
    add_bullet(doc, "Web app: Streamlit")
    add_bullet(doc, "Data processing: Python + Pandas")
    add_bullet(doc, "Database: SQLite และ CSV สำหรับ local/prototype")
    add_bullet(doc, "Visualization: Plotly")
    add_bullet(doc, "Testing: Pytest, validation scripts, privacy check, label validation และ Phase 8 acceptance script")
    add_bullet(doc, "Help AI: Ollama/localhost สำหรับตอบจากเอกสารโครงการเท่านั้น และยังใช้ได้แม้ไม่เปิด AI ผ่าน document search")
    add_bullet(doc, "Deployment option: Streamlit Community Cloud หรือ Hugging Face Spaces")

    add_heading(doc, "9. Data Pipeline และ Data Quality", 1)
    add_para(doc, "Pipeline ทำหน้าที่อ่าน workbook ตรวจ required sheets/columns ทำความสะอาด no-PII fields แปลง error เช่น #NUM! เป็น missing value parse date/income คำนวณ study duration และตรวจคุณภาพหลายมิติ ทั้ง duplicate ID, date order, income range, dictionary values, key-field completeness และ cross-field relationships")
    for item in [
        "Completeness Score วัดความครบถ้วนของข้อมูลสำคัญ",
        "Validity Score วัดความถูกต้องตามรูปแบบและช่วงค่าที่กำหนด",
        "Uniqueness Score ตรวจความซ้ำซ้อนของ identifier หรือ record",
        "Consistency Score ตรวจความสอดคล้องระหว่าง field เช่น graduation กับ employment follow-up",
        "Field cleaning report ระบุ readiness สำหรับ dashboard, analytics, policy และ ML พร้อม leakage risk",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "10. Risk, Forecast และ Policy Recommendation", 1)
    add_para(doc, "Prototype ใช้ rule-based risk score เป็นรุ่นแรกเพื่อให้ทุกคะแนนอธิบายได้และตรวจทานได้ ก่อนขยับสู่ ML ในอนาคต ระบบรายงาน total risk score, risk level, triggered components, component scores, calculation timestamp, rule version, expert approval status และ limitations")
    add_table(
        doc,
        ["ส่วนงาน", "หลักการ"],
        [
            ["Risk Score", "ใช้ rule/weight/threshold จาก config และรายงานเหตุผลของคะแนนทุกครั้ง"],
            ["Graduation Success", "จำแนก on-time, delayed, studying, over-duration risk, exited, unknown ด้วยกฎที่ตรวจสอบได้"],
            ["Policy Ranking", "จัดอันดับข้อเสนอจาก measured rates, records, formula, configurable weights และ evidence columns"],
            ["ข้อจำกัด", "ยังไม่ใช่ ML production และผลลัพธ์สนับสนุนการอภิปรายเชิงนโยบาย ไม่ใช่ automatic allocation"],
        ],
        [1.6, 4.9],
    )

    add_heading(doc, "11. Agents และ Human-in-the-loop Governance", 1)
    add_para(doc, "ระบบออกแบบ agent เป็น system modules เพื่อช่วยงานเฉพาะด้าน แต่การรับรอง label กฎความเสี่ยง น้ำหนักเชิงนโยบาย และข้อกำกับข้อมูลต้องทำโดยมนุษย์ที่รับผิดชอบ")
    add_table(
        doc,
        ["Agent", "บทบาท"],
        [
            ["Data Steward", "ตรวจคุณภาพข้อมูลและเสนอ cleaning issues"],
            ["Analytics", "สรุป descriptive/diagnostic analytics"],
            ["Risk Prediction", "คำนวณ risk score และอธิบายเหตุผล"],
            ["Policy Recommendation", "สร้างข้อเสนอเชิงนโยบายจากข้อมูลและน้ำหนักที่กำหนด"],
            ["External Indicator", "จัดการ template ข้อมูลเสริมรายปี"],
            ["Governance", "ดูแล PII masking, role concept และ audit/export log"],
            ["Technical", "ดูแล reproducibility, deployment, database, security, monitoring และ support"],
        ],
        [1.8, 4.7],
    )
    add_para(doc, "สถานะล่าสุดของ Prototype อนุญาตให้ใช้ Label ทั้ง 8 รายการสำหรับ prototype ML experiment only ตั้งแต่วันที่ 2026-07-28 แต่ยังไม่ใช่ production approval การใช้ผลกับบุคคลจริงหรือการจัดสรรทุนจริงยังต้องผ่านผู้เชี่ยวชาญและ governance gate")
    add_bullet(doc, "Label workflow ต้อง freeze source snapshot, เตรียม evidence, ตรวจ human approval, ตรวจ fairness/privacy และแยก train/test ตาม cohort หรือเวลา")
    add_bullet(doc, "Agent ช่วยเตรียม ตรวจ และจัดคิว review ได้ แต่ไม่ทดแทน Data owner, domain expert, Policy owner หรือ DPO")

    add_heading(doc, "12. Help, Documentation และ AI-assisted Integration", 1)
    add_para(doc, "ระบบล่าสุดมีหน้า Help & Documentation เพิ่มขึ้น เพื่อให้ผู้ใช้ค้นหาคู่มือ เรียนรู้ขั้นตอนสำคัญ อ่าน FAQ และถามผู้ช่วย AI ที่ทำงานภายในเครื่องจากเอกสารของโครงการ")
    add_table(
        doc,
        ["ส่วน", "ขอบเขต", "ข้อกำกับ"],
        [
            ["Documentation Search", "ค้นหาและเปิดเอกสารใน docs/*.md", "ใช้เฉพาะเอกสารโครงการ ไม่ดึงข้อมูลรายบุคคล"],
            ["FAQ", "ตอบคำถามใช้งาน เช่น target, forecast, privacy และ AI", "เป็นคำตอบเชิงคู่มือ ไม่ใช่คำตัดสินเชิงนโยบาย"],
            ["Local AI/Ollama", "ถามตอบจากเอกสารผ่าน endpoint localhost", "จำกัด localhost และไม่ส่งฐานข้อมูลผู้รับทุนหรือ PII"],
            ["Privacy Warning", "เตือนห้ามวางชื่อ เลขบัตร อีเมล เบอร์โทรศัพท์ หรือข้อมูลระบุตัวบุคคล", "หากเปลี่ยนไปใช้ AI ภายนอก ต้องผ่าน Privacy, DPA และนโยบายหน่วยงานก่อน"],
        ],
        [1.5, 3.0, 2.0],
    )

    add_heading(doc, "13. Privacy, Security และ Display Rules", 1)
    add_para(doc, "Prototype ใช้หลัก aggregate-first และ no-direct-PII โดยค่าเริ่มต้น ข้อมูลที่ห้ามแสดงหรือ export ได้แก่ ชื่อ-สกุลผู้ติดต่อ หมายเลขโทรศัพท์ เลขที่สัญญารับทุน เลขเอกสารรับรอง ที่อยู่ละเอียด และหมายเหตุที่อาจมีข้อมูลอ่อนไหว")
    add_bullet(doc, "แสดงข้อมูลรวมระดับ cohort, province, region, country, field group, status และ employment type")
    add_bullet(doc, "ใช้รหัสผู้รับทุนแบบไม่ระบุตัวตน เช่น odos_uid เมื่อจำเป็น")
    add_bullet(doc, "export เฉพาะข้อมูล aggregate และต้องตรวจสิทธิ์ตาม role mockup")
    add_bullet(doc, "มี privacy_check.py, governance/privacy.py และ test_privacy.py เป็นชั้นตรวจสอบใน Prototype")

    add_heading(doc, "14. เฟสงานและสถานะ", 1)
    add_table(
        doc,
        ["Phase", "เป้าหมาย", "สถานะ/ผลลัพธ์"],
        [
            ["0", "Lock prototype scope", "Locked: scope และ acceptance criteria"],
            ["1", "Prepare data", "พร้อมสำหรับ development รอ formal data definition sign-off"],
            ["2", "Prepare repository", "repository, AGENTS.md และ project structure พร้อม"],
            ["3", "Build app skeleton", "Streamlit app, navigation และ SQLite schema scaffolded"],
            ["4", "Build data pipeline", "import, cleaning, validation และ quality workflow implemented"],
            ["5", "Build dashboard/analytics", "Overview, Data Quality และ Analytics implemented"],
            ["6", "Risk/forecast/policy", "rule-based risk, graduation status และ policy ranking implemented"],
            ["7", "External indicators/governance", "template, PII masking, audit/export controls implemented"],
            ["8", "Test, deploy, handover", "functional, data, privacy, deployment, help และ handover acceptance checks implemented"],
            ["9", "Expand to production", "อนาคต: database server, auth, ML, API, PDPA workflow"],
        ],
        [0.7, 2.1, 3.7],
    )

    add_heading(doc, "15. Acceptance Criteria", 1)
    add_para(doc, "Prototype Phase 0-8 ถือว่าผ่านเมื่อระบบและเอกสารสามารถแสดงผลลัพธ์หลักได้ครบ พร้อมข้อจำกัดที่โปร่งใส")
    for item in [
        "มี scope ที่ล็อกขอบเขตและข้อจำกัดชัดเจน",
        "มีฐานข้อมูลกลางหรือ cleaned dataset ที่ใช้รันระบบได้",
        "มี web app ที่เปิดได้และมีหน้าหลักตาม navigation ล่าสุด รวม Help & Documentation",
        "Overview, Data Quality และ Analytics แสดง metric/filter/chart ได้ถูกต้องตามนิยาม",
        "Risk & Forecast มีคะแนนเบื้องต้น คำอธิบายที่มา rule version และ limitations",
        "Policy Recommendation แสดงข้อเสนอพร้อมเหตุผลจากข้อมูลและรองรับการปรับน้ำหนัก",
        "External Indicators มี template สำหรับข้อมูลรายปี และ Governance แสดง PII masking, role concept, audit/export log",
        "Help & Documentation ค้นหาเอกสารได้ และ Local AI จำกัดเฉพาะ localhost/เอกสารโครงการโดยไม่แตะฐานข้อมูลผู้รับทุน",
        "มี README คู่มือเบื้องต้น handover scenario acceptance scripts และ roadmap สู่ production",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "16. ข้อจำกัดและความเสี่ยงที่ต้องสื่อสาร", 1)
    for item in [
        "ระบบเป็น local SQLite/Streamlit Prototype ไม่ใช่ multi-user production service",
        "role selection เป็น mockup ยังไม่มี authentication/authorization จริง",
        "external indicators ยังเป็น template/sample ต้องมีแหล่งข้อมูลรายปีที่ตรวจสอบได้",
        "risk rules และ thresholds ต้องได้รับ expert approval ก่อนใช้งานเชิงปฏิบัติการ",
        "ยังไม่เปิดใช้ ML forecasting จนกว่าจะมีนิยาม outcome, train/test design, leakage control และ monitoring",
        "upload validation ยังไม่มี staging, approval workflow, rollback และ async processing ระดับ production",
        "Local AI เป็น optional helper สำหรับเอกสาร ไม่ใช่ระบบอนุมัติ ไม่ใช่โมเดลนโยบาย และต้องระวังไม่ป้อน PII ลงช่องถาม",
        "data quality warnings ต้องได้รับการแก้ไขหรือยอมรับโดย data owner ก่อนใช้ประกอบนโยบายจริง",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "17. ข้อมูลเพิ่มเติมที่ควรจัดหา", 1)
    for item in [
        "ข้อมูลความต้องการแรงงานรายปีตามสาขาและพื้นที่ พร้อมวันที่แหล่งข้อมูลและระดับความน่าเชื่อถือ",
        "ข้อมูลต้นทุนทุน ค่าใช้จ่าย และค่าครองชีพรายปลายทางหรือพื้นที่",
        "ตัวชี้วัดความเหลื่อมล้ำและการพัฒนาท้องถิ่นในระดับพื้นที่ที่ใช้วิเคราะห์",
        "ข้อมูลผลลัพธ์หลังจบทุนที่มีวันที่รายงาน ประเภทงาน กลุ่มอาชีพ รายได้ และสถานะกลับพื้นที่",
        "taxonomy ของพื้นที่และสาขาวิชาที่มี code, effective date และ mapping history",
        "กฎความเสี่ยง KPI thresholds recommendation weights และ sign-off records ที่ผ่านผู้เชี่ยวชาญ",
        "แนวทางใช้งาน AI ช่วยงานเอกสารในระดับองค์กร รวมถึง privacy review และ DPA หากใช้ endpoint ภายนอก",
        "ข้อกำหนด production security, audit retention, backup, recovery และ PDPA operating procedures",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "18. Roadmap สู่ Production", 1)
    add_para(doc, "การเข้าสู่ Phase 9 ควรเริ่มหลังจาก data owner รับรองนิยามและคุณภาพข้อมูล policy owner รับรองกฎและข้อเสนอ legal/governance owner อนุมัติแบบการใช้ข้อมูล และมีแผน hosting/security ที่ได้รับงบประมาณ")
    add_number(doc, "ยืนยัน data definition, label, KPI และ taxonomy ที่ใช้ร่วมกันทั้งโครงการ")
    add_number(doc, "ออกแบบ production architecture: database server, API, authentication, authorization, backup, monitoring และ deployment")
    add_number(doc, "จัดทำ PDPA workflow, data use notice, retention rules, access review และ incident response")
    add_number(doc, "ยกระดับ risk/forecast เป็น validated ML workflow เมื่อข้อมูลและ governance พร้อม")
    add_number(doc, "กำหนด AI assistance policy ว่าส่วนใดใช้ local-only ได้ ส่วนใดต้องห้าม และเงื่อนไขสำหรับ external AI")
    add_number(doc, "สร้าง operating model สำหรับ data refresh, model/rule review, policy sign-off และ change management")

    add_heading(doc, "19. แหล่งข้อมูลที่ใช้จัดทำเอกสาร", 1)
    for item in [
        "PROJECT_SCOPE.md และ odos-policy-analytics/PROJECT_SCOPE.md",
        "odos-policy-analytics/README.md",
        "odos-policy-analytics/docs/architecture.md",
        "odos-policy-analytics/docs/data_pipeline_quality.md",
        "odos-policy-analytics/docs/risk_forecast_policy.md",
        "odos-policy-analytics/docs/external_indicators_governance.md",
        "odos-policy-analytics/docs/agents_and_labeling.md",
        "odos-policy-analytics/docs/deployment.md",
        "odos-policy-analytics/docs/acceptance_tests.md",
        "odos-policy-analytics/docs/phase8_handover.md",
        "odos-policy-analytics/app.py และ pages/08_help_documentation.py",
        "odos-policy-analytics/config/app_config.yaml",
    ]:
        add_bullet(doc, item)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = footer.add_run("ODOS Policy Analytics Prototype | กรอบงานและขอบเขตโครงการ")
        set_run_font(r, 11, False, MUTED)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
