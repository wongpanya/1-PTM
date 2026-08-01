from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Policy_Data_AI_Training_Curriculum_Full_TH.docx"

FONT = "Arial"
NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F3F7FB"
GOLD = "C69214"
PALE_GOLD = "FFF8E8"
GRAY = "5E6873"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
BLACK = "000000"
GREEN = "2E7D5B"
RED = "A33A3A"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, size=11.5, bold=False, color=BLACK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph_runs(paragraph, size=11.5, bold=False, color=BLACK):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color)


def add_paragraph(doc, text="", bold_lead=None, style=None, after=6, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True, color=NAVY)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullets(doc, items, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(item)
        set_run_font(run)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(item)
        set_run_font(run)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        size, color, before, after = 16, BLUE, 18, 10
    elif level == 2:
        size, color, before, after = 13, BLUE, 14, 7
    else:
        size, color, before, after = 12, NAVY, 10, 5
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    style_paragraph_runs(p, size=size, bold=True, color=color)
    return p


def add_callout(doc, title, body, fill=PALE_GOLD, accent=GOLD):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11.5, bold=True, color=accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    set_run_font(r2, size=11)
    add_paragraph(doc, "", after=2)


def add_table(doc, headers, rows, widths, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=font_size, bold=True, color=WHITE)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(values):
            if row_idx % 2 == 1:
                set_cell_shading(cells[col_idx], PALE_BLUE)
            p = cells[col_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            if col_idx == 0 and len(headers) > 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value))
            set_run_font(run, size=font_size)
    add_paragraph(doc, "", after=2)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=GRAY)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in (
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, NAVY),
    ):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(11.5)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("คู่มือการอบรม | ข้อมูลและ AI เพื่อนโยบายการศึกษา")
    set_run_font(header_run, size=8.5, color=GRAY)
    add_page_number(section.footer.paragraphs[0])

    # Opening block: workshop_agenda pattern with a compact metric strip.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("คู่มือวิทยากรและแผนการอบรม")
    set_run_font(r, size=11, bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("การใช้ข้อมูลและ AI เพื่อออกแบบนโยบาย\nลดความเหลื่อมล้ำทางการศึกษา")
    set_run_font(r, size=25, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("หลักสูตรแกนกลางร่วมและห้องปฏิบัติการสำหรับโจทย์นโยบาย 4 กลุ่ม")
    set_run_font(r, size=13, color=GRAY)

    metric = doc.add_table(rows=1, cols=4)
    set_table_geometry(metric, [2340, 2340, 2340, 2340])
    metrics = [
        ("3 วัน", "ระยะเวลาแนะนำ"),
        ("18 ชั่วโมง", "เวลาเรียนรู้"),
        ("4 Tracks", "โจทย์เฉพาะกลุ่ม"),
        ("50%", "สัดส่วนลงมือทำ"),
    ]
    for idx, (value, label) in enumerate(metrics):
        cell = metric.cell(0, idx)
        set_cell_shading(cell, PALE_GOLD)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(1)
        rv = p1.add_run(value)
        set_run_font(rv, size=14, bold=True, color=GOLD)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        rl = p2.add_run(label)
        set_run_font(rl, size=8.5, color=GRAY)

    add_paragraph(doc, "", after=12)
    add_callout(
        doc,
        "หลักคิดสำคัญ",
        "หลักสูตรนี้ไม่ได้สอนให้ทุกกลุ่มใช้แบบจำลองเดียวกัน แต่สอนกระบวนการร่วมตั้งแต่การตั้งโจทย์ "
        "การตรวจข้อมูล การวิเคราะห์ความเหลื่อมล้ำ การประเมินความไม่แน่นอน "
        "จนถึงการออกแบบทางเลือกเชิงนโยบายที่ตรวจสอบได้",
    )

    add_paragraph(
        doc,
        "จัดทำจากการทบทวนโจทย์โครงการของทั้ง 4 กลุ่ม ข้อมูลตัวอย่าง Education_Inequality_Data.csv "
        "และกรอบการทำงานของระบบวิเคราะห์นโยบายต้นแบบ",
        after=2,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    p = doc.paragraphs[-1]
    style_paragraph_runs(p, size=9.5, color=GRAY)
    doc.add_page_break()

    add_heading(doc, "บทสรุปสำหรับผู้จัดการอบรม", 1)
    add_paragraph(
        doc,
        "กลุ่มผู้เข้าอบรมมีพื้นฐานการใช้คอมพิวเตอร์ อินเทอร์เน็ต และ AI ระดับต้น "
        "แต่ไม่มีความรู้ด้านการเขียนโปรแกรม ขณะเดียวกันโจทย์ทั้ง 4 กลุ่มมีธรรมชาติข้อมูลและวิธีวิเคราะห์ต่างกัน "
        "จึงควรจัดการเรียนรู้แบบ No-code, evidence-first และ project-based",
    )
    add_bullets(
        doc,
        [
            "ใช้ข้อมูลโรงเรียนตัวอย่างเป็นสนามฝึกร่วมในวันแรก เพื่อให้ทุกคนเข้าใจกระบวนการเดียวกัน",
            "ใช้ข้อมูลจริงหรือข้อมูลจำลองเฉพาะโครงการในวันที่สอง เพื่อฝึกวิธีวิเคราะห์ที่ตรงกับโจทย์",
            "ให้ทุกกลุ่มผลิตผลงานมาตรฐานชุดเดียวกัน ได้แก่ Problem Statement, Data Map, Analysis Plan, Dashboard Storyboard และ Policy Brief",
            "ประเมินความสามารถจากการใช้หลักฐานและการอธิบายข้อจำกัด มากกว่าความซับซ้อนของแบบจำลอง",
            "กำหนดให้ AI เป็นผู้ช่วยตั้งคำถาม สรุป และตรวจความครบถ้วน ไม่ใช่ผู้ตัดสินใจแทนผู้รับผิดชอบนโยบาย",
        ],
    )
    add_callout(
        doc,
        "ข้อเสนอเชิงบริหาร",
        "ระยะเวลา 3 วันเหมาะสมที่สุด หากลดเหลือ 2 วัน ควรตัดการสร้างแบบจำลองออก "
        "และให้แต่ละกลุ่มส่งเพียงแผนการวิเคราะห์กับต้นแบบข้อเสนอเชิงนโยบาย",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    add_heading(doc, "1. บริบทและโจทย์การออกแบบหลักสูตร", 1)
    add_paragraph(
        doc,
        "โจทย์ทั้งหมดเกี่ยวข้องกับการลดความเหลื่อมล้ำทางการศึกษา แต่ครอบคลุมตั้งแต่การบริหารกำลังครู "
        "การจัดสรรทุน การกำกับมาตรฐานปฐมวัย ไปจนถึงการพัฒนาสมรรถนะครูวิทยาศาสตร์ "
        "การอบรมจึงต้องรักษามาตรฐานกระบวนการร่วมโดยไม่บังคับให้ทุกกลุ่มใช้ข้อมูลหรือเครื่องมือเหมือนกัน",
    )
    add_table(
        doc,
        ["กลุ่ม", "โจทย์หลัก", "หน่วยวิเคราะห์", "ผลลัพธ์เป้าหมาย"],
        [
            ["ช่วยพี่หน่อย", "ความไม่สมดุลของกำลังครูภายใต้โครงสร้างประชากร", "จังหวัด โรงเรียน สาขาครู และปี", "แผนที่ความเสี่ยงและช่องว่างกำลังครู"],
            ["PTM", "ทุนการศึกษากับกำลังคนแห่งอนาคต", "ผู้รับทุน รุ่นทุน สาขา ประเทศ จังหวัด", "แนวทางจัดสรรทุนและติดตามความเสี่ยง"],
            ["พี่อ๋อและคณะ", "ความเสี่ยงของตัวบ่งชี้มาตรฐานปฐมวัย", "สถานพัฒนา ตัวบ่งชี้ พื้นที่ และปี", "Risk Index, Ranking และ Early Warning"],
            ["ดาวล้อมเดือน", "อนาคตการจัดการเรียนรู้ของครูวิทยาศาสตร์", "ครู โรงเรียน สมรรถนะ พื้นที่ และผลลัพธ์", "Gap Analysis, Sandbox และข้อเสนอพัฒนาครู"],
        ],
        [1500, 3120, 2220, 2520],
    )

    add_heading(doc, "2. หลักการออกแบบการเรียนรู้", 1)
    principles = [
        ("Common Core + Project Tracks", "ทุกกลุ่มเรียนแกนกลางร่วมกัน แล้วแยกคลินิกตามประเภทโจทย์"),
        ("No-code First", "ใช้ตาราง ตัวกรอง Dashboard และ AI โดยไม่เริ่มจากการเขียนโปรแกรม"),
        ("Evidence before Model", "เริ่มจากคุณภาพข้อมูลและสถานการณ์ฐานก่อนเลือกวิธีพยากรณ์"),
        ("Learning by Producing", "ทุกช่วงต้องสร้างชิ้นงานที่นำไปประกอบโครงการจริงได้"),
        ("Uncertainty is Visible", "ผลวิเคราะห์ทุกชิ้นต้องระบุสมมติฐาน ข้อจำกัด และสิ่งที่ยังไม่ทราบ"),
        ("Human-in-the-loop", "ข้อเสนอจากแบบจำลองหรือ AI ต้องผ่านการตรวจสอบโดยผู้รับผิดชอบ"),
        ("Equity and Privacy by Design", "คำนึงถึงความเป็นธรรม การไม่ตีตรา และการใช้ข้อมูลเท่าที่จำเป็น"),
    ]
    for title, body in principles:
        add_paragraph(doc, f"{title}: {body}", bold_lead=f"{title}:")

    add_heading(doc, "3. กลุ่มเป้าหมายและเงื่อนไขก่อนเข้าอบรม", 1)
    add_bullets(
        doc,
        [
            "ผู้บริหาร นักวิเคราะห์นโยบาย นักวิชาการ ผู้รับผิดชอบข้อมูล และผู้ประสานงานโครงการ",
            "สามารถใช้เว็บเบราว์เซอร์ โปรแกรมตารางคำนวณ และเครื่องมือ AI เบื้องต้น",
            "ไม่จำเป็นต้องมีพื้นฐานสถิติขั้นสูงหรือการเขียนโปรแกรม",
            "แต่ละกลุ่มควรมีสมาชิกอย่างน้อย 3 บทบาท ได้แก่ ผู้รู้โจทย์ ผู้รู้ข้อมูล และผู้สื่อสารนโยบาย",
            "นำข้อมูลจริงแบบไม่ระบุตัวบุคคล หรือใช้ข้อมูลจำลองที่มีโครงสร้างใกล้เคียงข้อมูลจริง",
        ],
    )

    add_heading(doc, "4. ผลลัพธ์การเรียนรู้", 1)
    add_numbers(
        doc,
        [
            "กำหนดปัญหา กลุ่มเป้าหมาย พื้นที่ ช่วงเวลา และผลลัพธ์เชิงนโยบายได้",
            "ระบุหน่วยวิเคราะห์ ตัวแปร แหล่งข้อมูล และความสัมพันธ์ระหว่างตารางข้อมูลได้",
            "ตรวจพบข้อมูลขาดหาย ซ้ำ ผิดช่วง ไม่สอดคล้อง หรือมีความเสี่ยงด้านความเป็นส่วนตัวได้",
            "เลือกการวิเคราะห์เชิงพรรณนา เปรียบเทียบ จัดกลุ่ม พยากรณ์ หรือฉากทัศน์ให้เหมาะกับโจทย์ได้",
            "อ่านกราฟ คะแนนความเสี่ยง และผลพยากรณ์โดยไม่ตีความเกินหลักฐานได้",
            "ใช้ AI ช่วยตั้งคำถาม สรุปผล และตรวจความครบถ้วน โดยไม่เปิดเผยข้อมูลอ่อนไหว",
            "พัฒนาทางเลือกเชิงนโยบายที่เชื่อมหลักฐาน กลุ่มเป้าหมาย กิจกรรม และตัวชี้วัดได้",
            "นำเสนอข้อเสนอพร้อมสมมติฐาน ข้อจำกัด และกลไกกำกับดูแลได้",
        ],
    )

    add_heading(doc, "5. สถาปัตยกรรมการเรียนรู้", 1)
    add_table(
        doc,
        ["ชั้นการเรียนรู้", "สัดส่วน", "สาระสำคัญ", "รูปแบบ"],
        [
            ["แกนกลาง", "60%", "โจทย์นโยบาย ข้อมูล คุณภาพข้อมูล ความเหลื่อมล้ำ การอ่านผล AI และธรรมาภิบาล", "บรรยายสั้น สาธิต และแบบฝึกข้อมูลกลาง"],
            ["Project Track", "25%", "วิธีวิเคราะห์และข้อจำกัดเฉพาะโครงการ", "คลินิกกลุ่มและลงมือทำ"],
            ["Policy Studio", "15%", "สังเคราะห์หลักฐาน ทางเลือก การดำเนินงาน และ KPI", "นำเสนอ รับคำวิจารณ์ และปรับผลงาน"],
        ],
        [1500, 900, 4020, 2940],
    )
    add_paragraph(
        doc,
        "ลำดับกระบวนการกลาง: โจทย์นโยบาย -> หน่วยวิเคราะห์และข้อมูล -> คุณภาพข้อมูล -> "
        "สถานการณ์ฐาน -> ความเหลื่อมล้ำและกลุ่มเสี่ยง -> พยากรณ์หรือฉากทัศน์ -> "
        "ทางเลือกเชิงนโยบาย -> KPI และธรรมาภิบาล",
        after=8,
    )

    doc.add_page_break()
    add_heading(doc, "6. ตารางการอบรม 3 วัน", 1)
    add_heading(doc, "วันที่ 1: จากปัญหานโยบายสู่หลักฐาน", 2)
    add_table(
        doc,
        ["เวลา", "หัวข้อ", "กิจกรรม", "ผลงาน"],
        [
            ["09.00-09.30", "เปิดการอบรมและ Pre-assessment", "สำรวจความคาดหวังและกรณีตัวอย่าง", "Baseline ผู้เข้าอบรม"],
            ["09.30-10.30", "ความเหลื่อมล้ำและโจทย์นโยบาย", "แยกอาการ สาเหตุ กลุ่มเป้าหมาย และผลลัพธ์", "Problem Statement"],
            ["10.45-12.00", "จากคำถามสู่ข้อมูล", "กำหนดหน่วยวิเคราะห์ ตัวแปร และช่วงเวลา", "Problem-to-Data Map"],
            ["13.00-14.15", "Data Literacy แบบไม่เขียนโปรแกรม", "อ่านข้อมูลโรงเรียน 500 แห่งและ Data Dictionary", "Variable Classification"],
            ["14.15-15.15", "Data Quality และ Privacy", "ตรวจค่าขาดหาย ซ้ำ ผิดช่วง และข้อมูลอ่อนไหว", "Quality Register"],
            ["15.30-16.30", "สถานการณ์ฐานและความเหลื่อมล้ำ", "เปรียบเทียบภูมิภาค พื้นที่ และกลุ่มความยากจน", "Evidence Log 3 ข้อ"],
        ],
        [1200, 2340, 3720, 2100],
    )

    add_heading(doc, "วันที่ 2: จากหลักฐานสู่การวิเคราะห์เฉพาะโครงการ", 2)
    add_table(
        doc,
        ["เวลา", "หัวข้อ", "กิจกรรม", "ผลงาน"],
        [
            ["09.00-10.00", "เลือกวิธีวิเคราะห์ให้ตรงคำถาม", "จำแนก Descriptive, Diagnostic, Predictive, Scenario และ Evaluation", "Method Selection"],
            ["10.00-11.00", "การพยากรณ์และความไม่แน่นอน", "Baseline, horizon, scenario, validation และช่วงความเชื่อมั่น", "Assumption Log"],
            ["11.15-12.00", "การใช้ AI สนับสนุนงานวิเคราะห์", "Prompt, fact-check และการป้องกันข้อมูลส่วนบุคคล", "AI Use Record"],
            ["13.00-15.00", "Project Track Clinic", "แยก 4 กลุ่มตามประเภทโจทย์", "Analysis Plan"],
            ["15.15-16.30", "ออกแบบ Dashboard Storyboard", "เลือก KPI กราฟ ตัวกรอง และข้อความเตือน", "Dashboard Sketch"],
        ],
        [1200, 2460, 3660, 2040],
    )

    add_heading(doc, "วันที่ 3: จากผลวิเคราะห์สู่การตัดสินใจเชิงนโยบาย", 2)
    add_table(
        doc,
        ["เวลา", "หัวข้อ", "กิจกรรม", "ผลงาน"],
        [
            ["09.00-10.00", "อ่านผลอย่างไม่ตีความเกินจริง", "แยกข้อเท็จจริง การอนุมาน ผลแบบจำลอง และข้อเสนอ", "Claim-Evidence Check"],
            ["10.00-11.15", "ออกแบบทางเลือกเชิงนโยบาย", "เปรียบเทียบผลกระทบ ความเป็นไปได้ ต้นทุน และความเสี่ยง", "Policy Options"],
            ["11.15-12.00", "KPI และแผนติดตามผล", "กำหนด Output, Outcome, baseline, target และแหล่งข้อมูล", "M&E Plan"],
            ["13.00-14.30", "Policy Studio", "จัดทำ Policy Brief และแผนดำเนินงาน", "ฉบับร่างกลุ่ม"],
            ["14.30-16.00", "นำเสนอกลุ่ม", "นำเสนอ 7 นาที ถามตอบ 5 นาที", "ข้อเสนอแนะจากคณะ"],
            ["16.00-16.30", "Post-assessment และแผนต่อยอด", "สะท้อนการเรียนรู้และกำหนดงานหลังอบรม", "Action Plan 30 วัน"],
        ],
        [1200, 2460, 3660, 2040],
    )

    add_heading(doc, "7. แผนการสอนรายโมดูล", 1)
    modules = [
        ("โมดูล 1: Policy Problem Framing", "เปลี่ยนหัวข้อกว้างให้เป็นคำถามที่ตัดสินใจได้", "Who-What-Where-When-Why, Theory of Change, กลุ่มเป้าหมาย", "เขียนโจทย์ 1 ย่อหน้าและคำถามวิเคราะห์ 3 ข้อ", "ระวังโจทย์ที่รวมหลายผลลัพธ์หรือหลายระดับนโยบายในครั้งเดียว"),
        ("โมดูล 2: Data Mapping", "มองเห็นว่าข้อมูลใดตอบคำถามใด", "หน่วยวิเคราะห์ คีย์เชื่อม ตารางข้อมูล มิติเวลา และภูมิศาสตร์", "สร้าง Data Map", "เน้นว่าข้อมูลระดับพื้นที่ตอบคำถามรายบุคคลไม่ได้"),
        ("โมดูล 3: Data Quality and Privacy", "ตัดสินความพร้อมของข้อมูลก่อนวิเคราะห์", "Completeness, validity, consistency, uniqueness, timeliness และ PII", "ตรวจข้อมูลตัวอย่างและบันทึกปัญหา", "ให้คะแนนความรุนแรงและผลต่อข้อสรุป"),
        ("โมดูล 4: Baseline and Inequality", "อธิบายสถานการณ์และความแตกต่างระหว่างกลุ่ม", "จำนวน สัดส่วน ค่าเฉลี่ย การกระจาย และการแบ่งกลุ่ม", "เขียนข้อค้นพบพร้อมตัวเลขสนับสนุน", "ห้ามใช้ค่าเฉลี่ยเพียงค่าเดียวแทนความหลากหลายทั้งหมด"),
        ("โมดูล 5: Method Selection", "เลือกวิธีให้เหมาะกับคำถามและข้อมูล", "Trend, risk, forecast, scenario, gap และ evaluation", "จับคู่คำถามกับวิธีวิเคราะห์", "แบบจำลองซับซ้อนไม่ได้ชดเชยข้อมูลที่สั้นหรือไม่ครบ"),
        ("โมดูล 6: AI-assisted Analysis", "ใช้ AI เพิ่มคุณภาพการคิดโดยไม่ลดการตรวจสอบ", "Prompt structure, verification, privacy และ audit trail", "ทดลอง Prompt และตรวจ hallucination", "ใช้เฉพาะข้อมูลสรุปหรือข้อมูลจำลองในบริการ AI ภายนอก"),
        ("โมดูล 7: Dashboard and Story", "ออกแบบหน้าจอเพื่อการตัดสินใจ", "KPI, comparison, trend, risk tier, filter และ caveat", "วาด Dashboard Storyboard", "หนึ่งหน้าควรตอบคำถามหลักหนึ่งชุด"),
        ("โมดูล 8: Policy Options", "เปลี่ยนข้อค้นพบเป็นทางเลือกที่เปรียบเทียบได้", "targeting, intervention, feasibility, cost, risk และ equity", "สร้างทางเลือกอย่างน้อย 3 ทางเลือก", "แยกสิ่งที่ข้อมูลบอกออกจากคุณค่าหรือดุลยพินิจ"),
        ("โมดูล 9: Monitoring and Governance", "กำหนดการติดตามและความรับผิดชอบ", "KPI, owner, review cycle, access control และ human review", "จัดทำ M&E และ Governance Checklist", "คะแนนความเสี่ยงไม่ควรถูกใช้เป็นคำตัดสินอัตโนมัติ"),
        ("โมดูล 10: Policy Communication", "นำเสนอหลักฐานแบบกระชับและตรวจสอบได้", "headline, evidence, uncertainty, recommendation และ next step", "นำเสนอ 7 นาที", "ให้ผู้ฟังเห็นทั้งสิ่งที่ทราบและสิ่งที่ยังไม่ทราบ"),
    ]
    for name, objective, content, activity, note in modules:
        add_heading(doc, name, 2)
        add_paragraph(doc, f"วัตถุประสงค์: {objective}", bold_lead="วัตถุประสงค์:")
        add_paragraph(doc, f"สาระสำคัญ: {content}", bold_lead="สาระสำคัญ:")
        add_paragraph(doc, f"กิจกรรม: {activity}", bold_lead="กิจกรรม:")
        add_paragraph(doc, f"ข้อสังเกตสำหรับวิทยากร: {note}", bold_lead="ข้อสังเกตสำหรับวิทยากร:")

    doc.add_page_break()
    add_heading(doc, "8. ห้องปฏิบัติการเฉพาะโครงการ", 1)

    add_heading(doc, "Track A: Thailand Teacher Foresight Map", 2)
    add_bullets(
        doc,
        [
            "คำถามหลัก: จังหวัดหรือสาขาวิชาใดมีแนวโน้มครูขาด ครูเกิน หรือเกิด Spatial Mismatch",
            "ข้อมูลขั้นต่ำ: นักเรียน ครูตามอายุและสาขา การเกษียณ การรับเข้า/ย้ายออก โรงเรียน พื้นที่ และประชากรวัยเรียน",
            "ลำดับวิเคราะห์: Baseline -> Stock-flow -> Demand-Supply Gap -> Scenario -> Risk Tier -> Map",
            "ฉากทัศน์อย่างน้อย 3 แบบ: แนวโน้มเดิม การเกษียณสูง และการจัดสรร/เคลื่อนย้ายเชิงนโยบาย",
            "ผลงาน: Gap table รายจังหวัด แผนที่ความเสี่ยง และมาตรการตามระดับ Ready/Monitor/Priority/Urgent",
        ],
    )
    add_callout(
        doc,
        "ข้อจำกัดสำคัญ",
        "ประวัติย้อนหลัง 5 ปีไม่เพียงพอสำหรับการพยากรณ์ 10 ปีด้วยการลากแนวโน้มอย่างเดียว "
        "ควรใช้แบบจำลองโครงสร้างกำลังคนและเปิดเผยสมมติฐานการเกิด การเกษียณ การย้าย และเกณฑ์อัตรากำลัง",
        fill="FDEEEE",
        accent=RED,
    )

    add_heading(doc, "Track B: ODOS / Future Workforce Policy", 2)
    add_bullets(
        doc,
        [
            "คำถามหลัก: ผู้รับทุนกลุ่มใดมีความเสี่ยง และควรจัดสรรทุนให้สาขา ประเทศ หรือพื้นที่ใด",
            "ข้อมูลขั้นต่ำ: รุ่นทุน พื้นที่ สาขา ประเทศ สถานะการศึกษา การสำเร็จ การทำงาน และความต้องการกำลังคน",
            "ลำดับวิเคราะห์: Cohort profile -> Outcome -> Risk segmentation -> Field/area mismatch -> Allocation scenario",
            "ขอบเขตต้นแบบแนะนำ: เลือก Scholarship Risk หรือ Scholarship Allocation เป็นแกนหลักหนึ่งเรื่อง",
            "ผลงาน: Cohort dashboard, risk profile, allocation options และพื้นที่ที่ควรเพิ่มโอกาส",
        ],
    )
    add_callout(
        doc,
        "ข้อจำกัดสำคัญ",
        "ROI, SROI และผลกระทบต่อผลิตภาพระดับประเทศต้องใช้ข้อมูลต้นทุน ผลลัพธ์ระยะยาว "
        "กลุ่มเปรียบเทียบ และสมมติฐานการให้คุณค่าที่ตรวจสอบได้ จึงควรจัดเป็นระยะต่อยอด",
        fill="FDEEEE",
        accent=RED,
    )

    add_heading(doc, "Track C: ECD Indicator Risk Forecast", 2)
    add_bullets(
        doc,
        [
            "คำถามหลัก: ตัวบ่งชี้ใดและพื้นที่ใดมีการไม่ผ่านต่อเนื่องหรือมีแนวโน้มแย่ลง",
            "ข้อมูลขั้นต่ำ: สถานพัฒนา พื้นที่ หน่วยงานกำกับ ตัวบ่งชี้ ผลผ่าน/ไม่ผ่าน และปีประเมิน",
            "ลำดับวิเคราะห์: Coverage -> Fail rate -> Trend -> Persistence -> Severity -> Composite risk -> Ranking",
            "Risk Index ควรแสดงสูตร น้ำหนัก และผลการทดสอบความไวเมื่อเปลี่ยนน้ำหนัก",
            "ผลงาน: Ranking, heat map, trend chart และรายการตัวบ่งชี้เร่งด่วน",
        ],
    )
    add_callout(
        doc,
        "ข้อจำกัดสำคัญ",
        "ข้อมูลรายปี พ.ศ. 2563-2568 มีเพียงประมาณ 6 จุดเวลา ซึ่งน้อยเกินไปสำหรับ ARIMA ที่มั่นคง "
        "ควรเริ่มจาก Trend, Persistence และ Scenario หรือเพิ่มข้อมูลรายเดือน/ไตรมาสก่อน",
        fill="FDEEEE",
        accent=RED,
    )

    add_heading(doc, "Track D: Science Teacher Futures and PISA", 2)
    add_bullets(
        doc,
        [
            "คำถามหลัก: สมรรถนะครูด้านใดมีช่องว่าง และมาตรการใดควรทดลองในบริบทโรงเรียนต่างกัน",
            "ข้อมูลขั้นต่ำ: ครู วิชาเอก ภาระงาน วิทยฐานะ การพัฒนา การจัดการเรียนรู้ บริบทโรงเรียน และผลลัพธ์ผู้เรียน",
            "ลำดับวิเคราะห์: Baseline -> Competency gap -> School segmentation -> Scenario -> Intervention design -> Sandbox",
            "แบ่งกลุ่มเป้าหมายตามความจำเป็นและศักยภาพ โดยหลีกเลี่ยงป้ายกำกับที่ตีตราโรงเรียน",
            "ผลงาน: competency gap map, intervention package, sandbox protocol และ KPI",
        ],
    )
    add_callout(
        doc,
        "ข้อจำกัดสำคัญ",
        "ไม่ควรอ้างว่าการพัฒนาครูทำให้คะแนน PISA เพิ่มขึ้นจากความสัมพันธ์เพียงอย่างเดียว "
        "ควรออกแบบการทดลองหรือ quasi-experiment และตรวจว่าข้อมูลครูกับผลลัพธ์ผู้เรียนเชื่อมโยงกันได้จริง",
        fill="FDEEEE",
        accent=RED,
    )

    add_heading(doc, "9. การใช้ข้อมูลกลาง Education_Inequality_Data.csv", 1)
    add_paragraph(
        doc,
        "ข้อมูลโรงเรียน 500 แห่ง 18 ตัวแปร เหมาะสำหรับสร้างภาษากลางก่อนแยก Track "
        "เพราะครอบคลุมบริบท ทรัพยากร ความยากจน และผลลัพธ์ แต่เป็นข้อมูลภาพตัดขวาง "
        "จึงใช้ฝึกการพยากรณ์ระยะยาวไม่ได้โดยตรง",
    )
    add_table(
        doc,
        ["ช่วงฝึก", "คำถาม", "ตัวแปรตัวอย่าง", "ทักษะ"],
        [
            ["สำรวจ", "ข้อมูลหนึ่งแถวหมายถึงอะไร", "school_id, province, school_type", "หน่วยวิเคราะห์และหมวดหมู่"],
            ["คุณภาพ", "ค่าผิดช่วงหรือคำจำกัดความใดไม่ชัด", "rates, score, computer_ratio", "Validation และ Data Dictionary"],
            ["เปรียบเทียบ", "กลุ่มใดมีผลลัพธ์ต่างกัน", "region, urban_rural, poverty_rate", "Filter, group และ benchmark"],
            ["ความสัมพันธ์", "ทรัพยากรและความเปราะบางสัมพันธ์กับผลลัพธ์อย่างไร", "budget, internet, O-NET, dropout", "Correlation ไม่เท่ากับ causation"],
            ["นโยบาย", "กลุ่มใดควรได้รับมาตรการแบบใด", "ตัวแปรหลายด้าน", "Targeting และ Policy Options"],
        ],
        [1320, 2700, 2700, 2640],
    )
    add_paragraph(
        doc,
        "ควรสร้างสำเนาสำหรับบท Data Quality ที่เพิ่มข้อมูลว่าง แถวซ้ำ ค่าผิดช่วง หมวดหมู่สะกดไม่ตรง "
        "และวันที่ไม่สอดคล้อง โดยเก็บไฟล์ต้นฉบับไว้เป็นชุดเฉลย",
    )

    add_heading(doc, "10. การเชื่อมกับระบบวิเคราะห์นโยบายต้นแบบ", 1)
    add_table(
        doc,
        ["ขั้นตอนการเรียนรู้", "หน้าระบบที่ใช้เป็นตัวอย่าง", "สิ่งที่ผู้เรียนต้องอธิบายได้"],
        [
            ["สถานการณ์ฐาน", "Overview", "KPI หลัก กลุ่มเปรียบเทียบ และขอบเขตข้อมูล"],
            ["ตรวจข้อมูล", "Data Quality", "ปัญหา ความรุนแรง และผลต่อการวิเคราะห์"],
            ["ค้นหารูปแบบ", "Analytics", "ข้อค้นพบพร้อมตัวเลขและตัวกรองที่ใช้"],
            ["ประเมินอนาคต", "Risk Forecast", "ความหมายของความเสี่ยง สมมติฐาน และความไม่แน่นอน"],
            ["ออกแบบมาตรการ", "Policy Recommendation", "เหตุผล กลุ่มเป้าหมาย ผลกระทบ และข้อจำกัด"],
            ["เพิ่มบริบท", "External Indicators", "ข้อมูลภายนอกที่ช่วยหรือเปลี่ยนการตีความ"],
            ["กำกับดูแล", "Governance", "ผู้มีสิทธิ์ใช้ข้อมูล การตรวจสอบ และ human review"],
            ["ขอความช่วยเหลือ", "Help & Documentation / Local AI", "การค้นเอกสารและใช้ AI โดยไม่ส่ง PII"],
        ],
        [2100, 2640, 4620],
    )

    add_heading(doc, "11. แนวทางใช้ AI ในห้องอบรม", 1)
    add_paragraph(
        doc,
        "โครงสร้าง Prompt กลาง: บทบาท + บริบท + ข้อมูลที่อนุญาต + งานที่ต้องการ + "
        "ข้อจำกัด + รูปแบบคำตอบ + การระบุสิ่งที่ยังไม่ทราบ",
    )
    add_table(
        doc,
        ["ใช้ AI ได้", "ต้องตรวจสอบ", "ไม่ควรทำ"],
        [
            ["ช่วยตั้งคำถามวิเคราะห์", "ตัวเลขทุกค่ากับข้อมูลต้นทาง", "ป้อนชื่อ เลขประจำตัว เบอร์โทร หรือข้อมูลอ่อนไหว"],
            ["ช่วยจัดหมวดตัวแปร", "ความหมายและหน่วยของตัวแปร", "ให้ AI ตัดสินสิทธิหรือจัดสรรโอกาสอัตโนมัติ"],
            ["ช่วยสรุปข้อค้นพบ", "ข้อความที่สรุปเหตุและผล", "ใช้คำตอบที่ไม่มีแหล่งหลักฐาน"],
            ["ช่วยเปรียบเทียบทางเลือก", "สมมติฐาน คุณค่า และผลกระทบต่อกลุ่มเปราะบาง", "ปกปิดว่าเนื้อหาส่วนใดมาจาก AI"],
        ],
        [3120, 3120, 3120],
    )

    add_heading(doc, "12. การประเมินผลการเรียนรู้", 1)
    add_table(
        doc,
        ["องค์ประกอบ", "น้ำหนัก", "หลักฐาน"],
        [
            ["ปัญหาและกลุ่มเป้าหมายชัดเจน", "15%", "Problem Statement และคำถามวิเคราะห์"],
            ["ใช้ข้อมูลและหลักฐานถูกต้อง", "20%", "Data Map, Evidence Log และกราฟ"],
            ["เข้าใจคุณภาพและข้อจำกัด", "15%", "Quality Register และ Assumption Log"],
            ["เลือกวิธีวิเคราะห์เหมาะสม", "15%", "Analysis Plan และเหตุผลการเลือก"],
            ["ข้อเสนอเชิงนโยบายสมเหตุสมผล", "20%", "Policy Options และ Policy Brief"],
            ["การดำเนินงานและ KPI", "10%", "Action Plan และ M&E Plan"],
            ["จริยธรรม ความเป็นธรรม และการสื่อสาร", "5%", "Governance Checklist และการนำเสนอ"],
        ],
        [4320, 1080, 3960],
    )
    add_callout(
        doc,
        "เกณฑ์ผ่าน",
        "คะแนนรวมไม่น้อยกว่า 70% และต้องไม่มีข้อผิดพลาดร้ายแรงด้านการเปิดเผยข้อมูลส่วนบุคคล "
        "การอ้างเหตุและผลโดยไม่มีหลักฐาน หรือการเสนอให้ระบบตัดสินใจแทนมนุษย์โดยไม่มีการทบทวน",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    add_heading(doc, "13. ผลงานที่แต่ละกลุ่มต้องส่ง", 1)
    add_numbers(
        doc,
        [
            "Project Question Canvas จำนวน 1 หน้า",
            "Data Map และ Data Dictionary ฉบับย่อ",
            "Data Quality and Limitation Register",
            "Analysis Plan พร้อมสมมติฐานและเกณฑ์ตรวจสอบ",
            "Evidence Log อย่างน้อย 3 ข้อ",
            "Dashboard Storyboard จำนวน 1-3 หน้า",
            "Policy Option Matrix อย่างน้อย 3 ทางเลือก",
            "Policy Brief จำนวน 1 หน้า",
            "แผนดำเนินงาน KPI และ Governance Checklist",
        ],
    )

    add_heading(doc, "14. ข้อมูลที่แต่ละกลุ่มควรเตรียมก่อนอบรม", 1)
    add_table(
        doc,
        ["Track", "ข้อมูลขั้นต่ำ", "ข้อมูลเสริม", "สิ่งที่ต้องยืนยัน"],
        [
            ["A กำลังครู", "ครู นักเรียน โรงเรียน อายุ สาขา พื้นที่ ปี", "การเกิด การย้าย เกณฑ์อัตรากำลัง", "นิยามครูขาด/เกินและ forecast horizon"],
            ["B ทุน", "รุ่นทุน พื้นที่ สาขา ประเทศ สถานะศึกษา/งาน", "ต้นทุน ความต้องการกำลังคน ผลลัพธ์ระยะยาว", "ผลลัพธ์หลักและขอบเขต PII"],
            ["C ปฐมวัย", "สถานพัฒนา ตัวบ่งชี้ ผลประเมิน พื้นที่ ปี", "ความรุนแรง การช่วยเหลือ การประเมินซ้ำ", "ความสม่ำเสมอของเกณฑ์แต่ละปี"],
            ["D ครูวิทยาศาสตร์", "ครู สมรรถนะ ภาระงาน โรงเรียน พื้นที่ ผลลัพธ์", "PLC การอบรม วิธีสอน บริบทผู้เรียน", "วิธีเชื่อมครูกับผลลัพธ์และแบบประเมินผล"],
        ],
        [1320, 3000, 2880, 2160],
    )

    add_heading(doc, "15. แผนเตรียมการของผู้จัด", 1)
    add_table(
        doc,
        ["ช่วงเวลา", "งาน"],
        [
            ["ก่อนอบรม 2-3 สัปดาห์", "ส่งแบบสำรวจพื้นฐาน รับโจทย์และ Data Inventory ตรวจสิทธิ์ใช้ข้อมูล"],
            ["ก่อนอบรม 1 สัปดาห์", "ตรวจไฟล์ เตรียมข้อมูลจำลอง สำรองอุปกรณ์ และจัด Track Clinic"],
            ["ก่อนอบรม 1 วัน", "ทดสอบระบบ บัญชีผู้ใช้ อินเทอร์เน็ต โปรเจกเตอร์ และไฟล์สำรองแบบออฟไลน์"],
            ["ระหว่างอบรม", "บันทึกคำถาม ประเด็นข้อมูล และสมมติฐานที่แต่ละกลุ่มต้องกลับไปตรวจ"],
            ["หลังอบรม 7 วัน", "รับฉบับปรับปรุงและให้คำแนะนำรายกลุ่ม"],
            ["หลังอบรม 30 วัน", "Project Clinic เพื่อตรวจข้อมูล วิธีวิเคราะห์ และแผนต้นแบบ"],
        ],
        [2280, 7080],
    )

    doc.add_page_break()
    add_heading(doc, "ภาคผนวก A: Project Question Canvas", 1)
    add_table(
        doc,
        ["หัวข้อ", "คำถามสำหรับกลุ่ม"],
        [
            ["ปัญหา", "ปัญหาที่ต้องตัดสินใจคืออะไร และใครได้รับผลกระทบ"],
            ["กลุ่มเป้าหมาย", "บุคคล โรงเรียน สถานพัฒนา พื้นที่ หรือระบบใด"],
            ["หน่วยวิเคราะห์", "หนึ่งแถวของข้อมูลควรหมายถึงอะไร"],
            ["ช่วงเวลา", "วิเคราะห์อดีตกี่ปี และมองอนาคตกี่ปี"],
            ["ผลลัพธ์", "ต้องการเปลี่ยนแปลงอะไร และวัดด้วยตัวชี้วัดใด"],
            ["การตัดสินใจ", "ใครจะใช้ผลวิเคราะห์เพื่อตัดสินใจเรื่องใด"],
            ["ข้อจำกัด", "ข้อมูลหรือปัจจัยใดที่ยังไม่มีและอาจเปลี่ยนข้อสรุป"],
        ],
        [2400, 6960],
    )

    add_heading(doc, "ภาคผนวก B: Data Readiness Checklist", 1)
    add_bullets(
        doc,
        [
            "มีเจ้าของข้อมูลและสิทธิ์ใช้งานชัดเจน",
            "มี Data Dictionary ระบุความหมาย หน่วย แหล่งที่มา และช่วงเวลา",
            "ระบุหน่วยวิเคราะห์และคีย์เชื่อมข้อมูลได้",
            "ตรวจข้อมูลขาดหาย ซ้ำ ค่าผิดช่วง และหมวดหมู่ไม่สอดคล้องแล้ว",
            "แยกข้อมูลส่วนบุคคลและกำหนดวิธีลดการระบุตัวตนแล้ว",
            "มีข้อมูลย้อนหลังเพียงพอกับวิธีพยากรณ์ที่เลือก",
            "นิยามผลลัพธ์และกลุ่มเป้าหมายไม่เปลี่ยนระหว่างปีโดยไม่บันทึก",
            "มีชุดข้อมูลสำหรับพัฒนาและชุดข้อมูลสำหรับตรวจสอบแยกกันเมื่อจำเป็น",
        ],
    )

    add_heading(doc, "ภาคผนวก C: Evidence Log", 1)
    add_table(
        doc,
        ["ข้อค้นพบ", "หลักฐาน/ตัวเลข", "กลุ่มหรือช่วงเวลา", "ข้อจำกัด", "ผลต่อนโยบาย"],
        [
            ["", "", "", "", ""],
            ["", "", "", "", ""],
            ["", "", "", "", ""],
        ],
        [1920, 1800, 1800, 1920, 1920],
    )

    add_heading(doc, "ภาคผนวก D: Policy Option Matrix", 1)
    add_table(
        doc,
        ["ทางเลือก", "กลุ่มเป้าหมาย", "ผลที่คาดหวัง", "ความเป็นไปได้", "ต้นทุน/ความเสี่ยง"],
        [
            ["ทางเลือกที่ 1", "", "", "", ""],
            ["ทางเลือกที่ 2", "", "", "", ""],
            ["ทางเลือกที่ 3", "", "", "", ""],
        ],
        [1560, 1920, 2040, 1800, 2040],
    )

    add_heading(doc, "ภาคผนวก E: โครงสร้างการนำเสนอ 7 นาที", 1)
    add_numbers(
        doc,
        [
            "ปัญหา กลุ่มเป้าหมาย และการตัดสินใจที่ต้องการสนับสนุน",
            "ข้อมูลที่ใช้ หน่วยวิเคราะห์ และข้อจำกัดสำคัญ",
            "ข้อค้นพบที่มีหลักฐานสนับสนุน 3 ข้อ",
            "ผลพยากรณ์ ฉากทัศน์ หรือการจัดลำดับความเสี่ยง",
            "ทางเลือกเชิงนโยบายและทางเลือกที่แนะนำ",
            "แผนดำเนินงาน KPI ความเสี่ยง และสิ่งที่ต้องตรวจสอบต่อ",
        ],
    )

    add_heading(doc, "ภาคผนวก F: แหล่งข้อมูลที่ใช้ทบทวน", 1)
    add_bullets(
        doc,
        [
            "ชื่อข้อเสนอโครงการ.docx: รายละเอียดโจทย์ของผู้เข้าอบรม 4 กลุ่ม",
            "Education_Inequality_Data.csv: ข้อมูลตัวอย่างระดับโรงเรียนสำหรับกิจกรรมแกนกลาง",
            "เอกสารกรอบงานและขอบเขตของระบบ ODOS Policy Analytics Prototype",
            "โครงสร้างหน้า Overview, Data Quality, Analytics, Risk Forecast, Policy Recommendation, External Indicators, Governance และ Help & Documentation",
        ],
    )

    add_callout(
        doc,
        "สถานะเอกสาร",
        "ฉบับนี้เป็นกรอบหลักสูตรและคู่มือวิทยากรตั้งต้น ควรปรับเวลา ตัวอย่าง และระดับความลึก "
        "หลังได้รับข้อมูลจริง จำนวนผู้เข้าอบรม และรายชื่อวิทยากรประจำ Track",
        fill=LIGHT_GRAY,
        accent=GRAY,
    )

    doc.core_properties.title = "แผนการสอนและอบรมการใช้ข้อมูลและ AI เพื่อนโยบายการศึกษา"
    doc.core_properties.subject = "หลักสูตรแกนกลางและ Project Track สำหรับโจทย์ลดความเหลื่อมล้ำทางการศึกษา"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "education policy, inequality, data analytics, AI, training"
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
